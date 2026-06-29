import os
import fitz  # PyMuPDF
import docx

def load_pdf(path: str) -> list[dict]:
    """
    Reads a PDF file into memory using PyMuPDF.
    Returns [{"page_number": int, "raw_text": str}] per page.
    """
    pages = []
    doc = fitz.open(path)
    for i, page in enumerate(doc):
        text = page.get_text()
        if text and text.strip():
            pages.append({
                "page_number": i + 1,
                "raw_text": text
            })
    return pages

def load_docx(path: str) -> list[dict]:
    """
    Reads a DOCX file into memory.
    Extracts text from both standard paragraphs and tables.
    Since docx has no native "pages," treat the whole doc as page_number = None.
    """
    doc = docx.Document(path)
    full_text = []
    
    # Python-docx document elements can be paragraphs or tables.
    # To preserve order perfectly is complex in python-docx, but iterating
    # paragraphs then tables is often sufficient, or iterating block elements.
    # We will extract paragraphs first, then tables.
    
    # 1. Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
            
    # 2. Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                if cell.text.strip():
                    row_data.append(cell.text.strip().replace("\n", " "))
            if row_data:
                full_text.append(" | ".join(row_data))
    
    text = "\n".join(full_text)
    
    return [{
        "page_number": None,
        "raw_text": text
    }]

def load_txt(path: str) -> list[dict]:
    """
    Reads a TXT file into memory.
    Treats the whole document as a single page.
    """
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
        
    return [{
        "page_number": None,
        "raw_text": text
    }]

def load_document(path: str) -> list[dict]:
    """
    Dispatcher that picks the right loader by extension.
    """
    _, ext = os.path.splitext(path)
    ext = ext.lower()
    
    if ext == ".pdf":
        return load_pdf(path)
    elif ext == ".docx":
        return load_docx(path)
    elif ext == ".txt":
        return load_txt(path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
